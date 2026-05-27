"""PRD endpoints with real AI generation"""

import asyncio
import logging
import base64
import difflib
import json
import os
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Optional, List, Dict, Any
from fastapi import APIRouter, Depends, Query, BackgroundTasks
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc, func

from app.core.html_sanitizer import sanitize_html


from app.core.database import get_db
from app.core.rate_limit import rate_limit
from app.core.responses import ResponseBuilder
from app.core.security import get_current_user_id
from app.services.ai_service import ai_service
from app.models.prd import PRD, PRDStatus
from app.models.prd_version import PRDVersion
from app.models.project import Project
from app.models.persona import Persona
from app.models.competitor import Competitor
from app.core.exceptions import AppException

router = APIRouter()


def prd_json_to_markdown(data: Dict[str, Any], title: str) -> str:
    """Convert AI-generated PRD JSON to markdown"""
    md = [f"# {title}"]

    outline = data.get("outline", {}) if isinstance(data.get("outline"), dict) else {}
    content = data.get("content", {}) if isinstance(data.get("content"), dict) else {}
    suggestions = data.get("suggestions", []) if isinstance(data.get("suggestions"), list) else []

    # Background
    background = content.get("background", {}) if isinstance(content, dict) else {}
    if isinstance(background, str):
        md.append("\n## 一、背景与目标")
        md.append(f"\n{background}")
    elif isinstance(background, dict) and background:
        md.append("\n## 一、背景与目标")
        md.append(f"\n### 执行摘要\n{background.get('executive_summary', '')}")

        bp = background.get("business_problem", {}) if isinstance(background.get("business_problem"), dict) else {}
        if bp:
            md.append("\n### 当前痛点")
            for pain in bp.get("pain_points", []) if isinstance(bp.get("pain_points"), list) else []:
                md.append(f"- {pain}")
            md.append(f"\n**现状流程**: {bp.get('current_state', '')}")

        md.append(f"\n### 产品愿景\n{background.get('product_vision', '')}")

        obj = background.get("objectives", {}) if isinstance(background.get("objectives"), dict) else {}
        if obj:
            md.append("\n### 项目目标")
            md.append("**主要目标**:")
            for o in obj.get("primary", []) if isinstance(obj.get("primary"), list) else []:
                md.append(f"- {o}")
            if obj.get("secondary"):
                md.append("\n**次要目标**:")
                for o in obj.get("secondary", []) if isinstance(obj.get("secondary"), list) else []:
                    md.append(f"- {o}")

        sm = background.get("success_metrics", {}) if isinstance(background.get("success_metrics"), dict) else {}
        if sm:
            md.append("\n### 成功指标")
            for k, v in sm.items():
                md.append(f"- {k}: {v}")

        scope = background.get("scope", {}) if isinstance(background.get("scope"), dict) else {}
        if scope:
            md.append("\n### 项目范围")
            md.append("**包含范围**:")
            for s in scope.get("in_scope", []) if isinstance(scope.get("in_scope"), list) else []:
                md.append(f"- {s}")
            md.append("\n**不包含范围**:")
            for s in scope.get("out_of_scope", []) if isinstance(scope.get("out_of_scope"), list) else []:
                md.append(f"- {s}")

    # User Stories
    user_stories = content.get("user_stories", []) if isinstance(content, dict) and isinstance(content.get("user_stories"), list) else []
    if user_stories:
        md.append("\n## 二、用户故事")
        for us in user_stories:
            if isinstance(us, dict):
                md.append(f"\n### {us.get('id', 'US-XXX')}: {us.get('role', '')}")
                md.append(f"\n> {us.get('story', '')}")
                md.append(f"\n**优先级**: {us.get('priority', 'P1')}")
                ac = us.get("acceptance_criteria", []) if isinstance(us.get("acceptance_criteria"), list) else []
                if ac:
                    md.append("\n**验收标准**:")
                    for c in ac:
                        md.append(f"- {c}")

    # Suggestions
    if suggestions:
        md.append("\n## 系统建议（AI生成）")
        for i, s in enumerate(suggestions, 1):
            md.append(f"{i}. {s}")

    # Outline
    sections = outline.get("sections", []) if isinstance(outline.get("sections"), list) else []
    if sections:
        md.append("\n---")
        md.append("\n## PRD标准结构")
        for sec in sections:
            if isinstance(sec, dict):
                md.append(f"{sec.get('chapter', '')}. {sec.get('title', '')}")
                for kp in sec.get("key_points", []) if isinstance(sec.get("key_points"), list) else []:
                    md.append(f"  - {kp}")

    return "\n".join(md)


def extract_chapter_content(markdown: str, chapter_titles: List[str]) -> Dict[str, str]:
    """按章节标题拆分 markdown，提取各章节正文"""
    contents: Dict[str, str] = {}
    if not markdown or not chapter_titles:
        return contents

    # 策略1: 匹配 markdown 标题 (# ~ ####)
    escaped_titles = [re.escape(t) for t in chapter_titles]
    pattern = r'(?:^|\n)(?:#{1,4}\s*)(' + '|'.join(escaped_titles) + r')\s*(?:\n|\r\n)'

    parts = re.split(pattern, markdown)
    current_title = None
    for part in parts:
        if part in chapter_titles:
            current_title = part
        elif current_title is not None:
            contents[current_title] = part.strip()
            current_title = None

    # 策略2: 如果策略1未命中，尝试匹配加粗标题或下划线标题
    if not contents:
        bold_pattern = r'(?:^|\n)\*\*(' + '|'.join(escaped_titles) + r')\*\*\s*(?:\n|\r\n)'
        parts = re.split(bold_pattern, markdown)
        for part in parts:
            if part in chapter_titles:
                current_title = part
            elif current_title is not None:
                contents[current_title] = part.strip()
                current_title = None

    # 策略3: 如果仍未命中，按通用章节关键词（如“背景与目标”、“用户故事”）做模糊拆分
    if not contents:
        generic_sections = ["背景与目标", "用户故事", "功能需求", "非功能需求", "产品概述"]
        for title in generic_sections:
            if title in markdown:
                idx = markdown.index(title)
                next_idx = len(markdown)
                for other in generic_sections:
                    if other != title and other in markdown:
                        o_idx = markdown.index(other)
                        if o_idx > idx and o_idx < next_idx:
                            next_idx = o_idx
                contents[title] = markdown[idx + len(title):next_idx].strip("\n#* ")

    return contents


class PRDCreateRequest(BaseModel):
    """Create PRD request"""
    project_id: str
    title: str = Field(..., min_length=1, max_length=200)
    template: Optional[str] = "standard"


class PRDUpdateRequest(BaseModel):
    """Update PRD request"""
    title: Optional[str] = Field(None, min_length=1, max_length=200)
    content: Optional[dict] = None
    status: Optional[str] = None
    markdown: Optional[str] = None


class PRDExportRequest(BaseModel):
    """Export PRD request"""
    format: str = Field("markdown", pattern="^(markdown|json|pdf)$")


class PRDGenerateRequest(BaseModel):
    """Generate PRD chapter request"""
    chapter: str
    prompt: str
    context: Optional[dict] = None
    bypass_cache: bool = False


@rate_limit(requests=100, window=60)
@router.get("", response_model=dict)
async def list_prds(
    project_id: Optional[str] = None,
    limit: int = 20,
    offset: int = 0,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """List PRDs with pagination"""
    base_query = select(PRD).where(PRD.created_by == user_id, PRD.deleted_at.is_(None))
    if project_id:
        base_query = base_query.where(PRD.project_id == project_id)
    base_query = base_query.order_by(desc(PRD.created_at))

    # Count total
    count_query = select(func.count()).select_from(base_query.subquery())
    total_result = await db.execute(count_query)
    total = total_result.scalar() or 0

    # Paginate
    query = base_query.offset(offset).limit(limit)
    result = await db.execute(query)
    prds = result.scalars().all()

    items = []
    for prd in prds:
        ai_gen = prd.ai_generated or {}
        doc_type = ai_gen.get("doc_type", "prd") if isinstance(ai_gen, dict) else "prd"
        items.append({
            "id": prd.id,
            "project_id": prd.project_id,
            "title": prd.title,
            "version": prd.version,
            "status": prd.status.value,
            "doc_type": doc_type,
            "created_at": prd.created_at.isoformat() if prd.created_at else None,
            "updated_at": prd.updated_at.isoformat() if prd.updated_at else None,
        })

    return ResponseBuilder.paginated(
        data=items,
        page=offset // limit + 1 if limit > 0 else 1,
        limit=limit,
        total=total
    )


# Template chapter definitions — differentiated structure per template
# Each template adds/removes chapters to match industry needs
TEMPLATE_CHAPTERS = {
    "default": [
        "背景与目标",
        "用户故事",
        "业务流程",
        "功能规格",
        "数据架构",
        "合规要求",
        "数据埋点",
        "里程碑",
    ],
    "medical": [
        "背景与目标",
        "用户故事",
        "业务流程",
        "功能规格",
        "数据架构",
        "合规要求",
        "多院区适配",
        "数据埋点",
        "里程碑",
    ],
    "saas": [
        "背景与目标",
        "用户故事",
        "业务流程",
        "功能规格",
        "数据架构",
        "合规要求",
        "租户与计费模型",
        "数据埋点",
        "里程碑",
    ],
    "ecommerce": [
        "背景与目标",
        "用户故事",
        "业务流程",
        "功能规格",
        "数据架构",
        "合规要求",
        "供应链与促销策略",
        "数据埋点",
        "里程碑",
    ],
}


# Template → industry mapping for consistent AI generation
TEMPLATE_TO_INDUSTRY = {
    "medical": "medical",
    "saas": "saas",
    "ecommerce": "ecommerce",
    "default": "general",
}

# Industry-specific framework infrastructure injected into initial PRD markdown.
# Each template provides skeleton tables, checklists, and templates for key chapters.
TEMPLATE_FRAMEWORK = {
    "default": {
        "1": """\n\n### 1.0 通用产品 MVP 范围决策矩阵\n\n| 产品形态 | 适用场景 | 包含章节 | 裁剪建议 |\n|:---|:---|:---|:---|\n| **单模块** | 只做一个功能模块 | 1,2,4,5,6 | 跳过3（简化流程）、7（基础埋点即可） |\n| **单系统** | 核心系统升级/替换 | 全部9章 | 标准流程 |\n| **平台型** | 数据中台/协同平台 | 全部9章 | 第5章重点扩展架构 |\n| **互联网** | 线上服务/内容平台 | 全部9章 | 第6章增加数据隐私法规 |\n\n#### 数据来源指引\n\n| 占位符类型 | 建议数据来源 | 获取方式 |\n|:---|:---|:---|\n| 性能指标 | 现有系统监控、APM工具 | 技术团队提供 |\n| 用户量/频次 | 产品分析系统统计 | 数据团队提供 |\n| 合规状态 | 最近一次安全/隐私评估报告 | 安全团队提供 |\n| 满意度 | 用户调研、NPS评分 | 用户研究团队 |\n| 成本数据 | 年度预算执行表、人力成本 | 财务/HR团队 |\n\n#### RACI 矩阵模板\n\n| 事项 | 负责(R) | 批准(A) | 咨询(C) | 告知(I) |\n|:---|:---|:---|:---|:---|\n| PRD评审 | 产品经理 | 项目负责人 | 技术负责人 | 测试、设计 |\n| 范围变更 | 产品经理 | 管理层 | 技术负责人 | 所有干系人 |\n| 合规审查 | 安全合规专员 | 安全负责人 | 法务 | 项目组 |\n| 上线决策 | 技术负责人 | 项目负责人 | 业务方 | 相关业务方 |\n""",
        "2": """\n\n### 2.0 通用用户故事验收标准模板\n\n每个用户故事必须包含：\n- **触发条件**：什么场景下触发此操作\n- **前置条件**：系统/用户必须满足的状态\n- **基本流程**：标准操作步骤，含系统反馈\n- **异常流程**：网络中断、权限不足、数据冲突等处理方式\n- **验收标准（AC）**：可测试的通过/失败条件\n- **优先级**：P0（必须）/ P1（应该）/ P2（可以）\n- **数据影响**：是否涉及敏感数据、是否需要审计留痕\n""",
        "5": """\n\n### 5.0 通用数据架构设计检查清单\n\n- [ ] 所有实体都有主键和业务唯一键\n- [ ] 敏感字段已标注加密/脱敏策略\n- [ ] 外键关系已定义级联规则\n- [ ] 大数据量表已规划分区/分表策略\n- [ ] 审计字段（created_at, updated_at, created_by）已覆盖全部操作\n- [ ] 归档和销毁策略已定义\n- [ ] 数据同步策略已明确\n- [ ] 接口数据交换格式已定义\n""",
        "6": """\n\n### 6.0 通用合规审查清单\n\n| 检查项 | 适用法规/标准 | 验证方式 | 责任人 |\n|:---|:---|:---|:---|\n| 网络安全与数据 | 网络安全法、数据安全法 | 安全评估报告 | 安全合规专员 |\n| 个人信息保护 | 个人信息保护法 | 隐私政策+用户授权记录 | 法务+产品经理 |\n| 数据分类分级 | 个保法、数据安全法 | 资产清单+分级矩阵 | 数据管理员 |\n| 操作审计留痕 | 等保要求 | 审计日志抽查 | 安全合规专员 |\n| 密码应用合规 | 密码法 | 安全评估报告 | 安全合规专员 |\n""",
        "8": """\n\n### 8.0 通用里程碑评审闸门（Stage-Gate）\n\n| 阶段 | 评审标准 | 通过条件 | 参与角色 |\n|:---|:---|:---|:---|\n| **Gate 0: 立项** | PRD框架通过评审 | 8章结构完整、范围明确 | PM、项目负责人 |\n| **Gate 1: 设计** | 详细设计稿+技术方案 | UI评审通过、架构评审通过 | PM、UI、架构师 |\n| **Gate 2: 开发** | 核心功能开发完成 | 单元测试覆盖率≥80% | 开发、测试 |\n| **Gate 3: 测试** | 集成测试+UAT通过 | 无P0/P1阻塞缺陷 | 测试、业务方 |\n| **Gate 4: 上线** | 灰度发布验证通过 | 核心指标无异常波动 | 运维、PM、业务方 |\n| **Gate 5: 运营** | 发布后30天评估 | 用户满意度≥目标值 | PM、运营、业务方 |\n""",
    },
    "medical": {
        "1": """\n\n### 1.0 医疗产品 MVP 范围决策矩阵\n\n| 产品形态 | 适用场景 | 包含章节 | 裁剪建议 |\n|:---|:---|:---|:---|\n| **单模块** | EMPI、权限中台等独立模块 | 1,2,4,5,6 | 跳过3（简化流程）、7（基础埋点即可） |\n| **单院区系统** | HIS/EMR/护理系统升级替换 | 全部9章 | 标准医疗流程 |\n| **多院区平台** | 区域医疗/医联体/数据中台 | 全部9章 | 第5章重点扩展联邦架构、第6章增加跨省法规 |\n| **互联网医院** | 线上诊疗/慢病管理/互联网+护理 | 全部9章 | 第6章增加互联网诊疗法规（《互联网诊疗监管细则》） |\n\n#### 医疗数据来源指引\n\n| 占位符类型 | 建议数据来源 | 获取方式 |\n|:---|:---|:---|\n| 性能指标 | 现有系统慢查询日志、APM监控 | 信息科提供 |\n| 用户量/频次 | HIS系统使用统计、CA登录日志 | 信息科提供 |\n| 合规状态 | 最近一次等保/密评/互联互通报告 | 安全科提供 |\n| 满意度 | 年度患者满意度调查、医护系统评价 | 质控科提供 |\n| 成本数据 | 年度IT预算执行表、厂商合同 | 财务科提供 |\n\n#### RACI 矩阵模板\n\n| 事项 | 负责(R) | 批准(A) | 咨询(C) | 告知(I) |\n|:---|:---|:---|:---|:---|\n| PRD评审 | 产品经理 | 项目负责人 | 技术负责人、医务科 | 信息科、测试 |\n| 范围变更 | 产品经理 | 医院管理层 | 技术负责人 | 所有干系人 |\n| 合规审查 | 安全合规专员 | 信息科主任 | 法务、等保机构 | 项目组 |\n| 上线决策 | 技术负责人 | 院长/分管副院长 | 医务科、护理部 | 全院科室 |\n""",
        "2": """\n\n### 2.0 医疗用户故事验收标准模板\n\n每个用户故事必须包含：\n- **触发条件**：什么临床场景下触发（如：患者挂号后、检验结果异常时）\n- **前置条件**：系统/用户必须满足的状态（如：医生已登录、患者已建档）\n- **基本流程**：标准操作步骤，含系统反馈\n- **异常流程**：网络中断、权限不足、数据冲突等处理方式\n- **验收标准（AC）**：可测试的通过/失败条件，需覆盖等保审计要求\n- **优先级**：P0（必须）/ P1（应该）/ P2（可以）\n- **合规备注**：是否涉及敏感数据、是否需要双签名、是否留痕\n""",
        "5": """\n\n### 5.0 医疗数据架构设计检查清单\n\n- [ ] 所有实体都有主键和业务唯一键（如患者ID、就诊流水号）\n- [ ] 敏感字段已标注加密/脱敏策略（姓名、身份证号、手机号、诊断结果）\n- [ ] 外键关系已定义级联规则（患者删除时的病历归档策略）\n- [ ] 大数据量表已规划分区/分表策略（检验结果、操作日志按时间分区）\n- [ ] 审计字段（created_at, updated_at, created_by）已覆盖全部医疗操作\n- [ ] 归档和销毁策略已定义（符合《医疗机构病历管理规定》留存期）\n- [ ] 多院区场景下的数据同步策略已明确（主从/联邦/湖仓架构选型）\n- [ ] 危急值数据已配置实时推送通道\n""",
        "6": """\n\n### 6.0 医疗合规审查清单\n\n| 检查项 | 适用法规 | 验证方式 | 责任人 |\n|:---|:---|:---|:---|\n| 等保三级 | 网络安全等级保护2.0 | 等保测评报告 | 安全合规专员 |\n| 数据分类分级 | 个人信息保护法、数据安全法 | 数据资产清单+分级矩阵 | 数据管理员 |\n| 患者知情同意 | 电子病历应用管理规范 | 系统功能截图+流程记录 | 产品经理 |\n| 操作审计留痕 | 等保2.0三级要求 | 审计日志样本抽查 | 安全合规专员 |\n| 跨境数据传输 | 数据出境安全评估办法 | 数据流向图+合规声明 | 法务+安全 |\n| 密码应用合规 | 密码法、GM/T 0054 | 密评报告 | 安全合规专员 |\n| 互联网诊疗合规 | 《互联网诊疗监管细则》 | 执业许可+流程留痕 | 医务科+法务 |\n""",
        "8": """\n\n### 8.0 医疗里程碑评审闸门（Stage-Gate）\n\n| 阶段 | 评审标准 | 通过条件 | 参与角色 |\n|:---|:---|:---|:---|\n| **Gate 0: 立项** | PRD框架通过评审 | 8章结构完整、范围明确 | PM、项目负责人 |\n| **Gate 1: 设计** | 详细设计稿+技术方案 | UI评审通过、架构评审通过 | PM、UI、架构师 |\n| **Gate 2: 开发** | 核心功能开发完成 | 单元测试覆盖率≥80% | 开发、测试 |\n| **Gate 3: 测试** | 集成测试+UAT通过 | 无P0/P1阻塞缺陷 | 测试、业务方 |\n| **Gate 3.5: 等保测评** | 等保三级测评通过 | 测评报告无高风险项 | 安全合规专员 |\n| **Gate 4: 上线** | 灰度发布验证通过 | 核心指标无异常波动 | 运维、PM、业务方 |\n| **Gate 5: 运营** | 发布后30天评估 | 用户满意度≥目标值、零重大安全事件 | PM、运营、业务方 |\n""",
    },
    "saas": {
        "1": """\n\n### 1.0 SaaS 产品 MVP 范围决策矩阵\n\n| 产品形态 | 适用场景 | 包含章节 | 裁剪建议 |\n|:---|:---|:---|:---|\n| **单租户版** | 大客户私有化部署 | 1,2,4,5,6 | 跳过3（定制流程）、7（基础埋点） |\n| **多租户SaaS** | 标准化SaaS订阅服务 | 全部9章 | 第5章重点扩展租户隔离架构 |\n| **API平台** | 面向开发者的能力开放 | 全部9章 | 第4章增加API设计规范、第6章增加OAuth合规 |\n| **白标/OEM** | 渠道/代理商贴牌销售 | 全部9章 | 第4章增加品牌化配置、第6章增加渠道协议 |\n\n#### SaaS 关键指标数据来源\n\n| 占位符类型 | 建议数据来源 | 获取方式 |\n|:---|:---|:---|\n| 注册转化率 | 产品分析系统（Mixpanel/Amplitude） | 增长团队 |\n| 活跃用户数 | 数据库活跃会话统计 | 数据团队 |\n| MRR/ARR | 计费系统/财务系统 | 财务团队 |\n| NPS | 用户满意度调研 | 客户成功团队 |\n| 支持工单量 | 客服系统（Zendesk/Intercom） | 客户成功团队 |\n\n#### RACI 矩阵模板\n\n| 事项 | 负责(R) | 批准(A) | 咨询(C) | 告知(I) |\n|:---|:---|:---|:---|:---|\n| PRD评审 | 产品经理 | 产品总监 | 技术负责人、设计师 | 销售、客户成功 |\n| 范围变更 | 产品经理 | 产品总监 | 技术负责人 | 所有干系人 |\n| 安全审查 | 安全工程师 | CTO | 法务 | 客户成功 |\n| 发布决策 | 技术负责人 | 产品总监 | 客户成功 | 所有客户（公告） |\n""",
        "2": """\n\n### 2.0 SaaS 用户故事验收标准模板\n\n每个用户故事必须包含：\n- **触发条件**：什么场景下触发（如：新租户注册后、订阅即将到期时）\n- **前置条件**：系统/用户必须满足的状态（如：租户已创建、管理员权限已分配）\n- **基本流程**：标准操作步骤，含多端响应（Web/App/小程序）\n- **异常流程**：租户隔离冲突、配额超限、权限不足等处理方式\n- **验收标准（AC）**：可测试的通过/失败条件，需覆盖多租户场景\n- **优先级**：P0（必须）/ P1（应该）/ P2（可以）\n- **计费影响**：是否影响MRR、是否涉及套餐升级/降级\n""",
        "5": """\n\n### 5.0 SaaS 数据架构设计检查清单\n\n- [ ] 租户隔离策略已明确（独立数据库/共享数据库+租户ID行隔离/Schema隔离）\n- [ ] 敏感字段已标注加密/脱敏策略（PII、支付信息、企业数据）\n- [ ] 外键关系已定义级联规则（租户删除时的数据清理策略）\n- [ ] 大数据量表已规划分区/分表策略（事件日志、操作记录按租户+时间分区）\n- [ ] 审计字段（created_at, updated_at, created_by, tenant_id）已覆盖全部操作\n- [ ] 数据保留策略已定义（符合GDPR/CCPA删除权要求）\n- [ ] 备份恢复策略已明确（RTO/RPO目标、跨地域容灾）\n- [ ] 多租户性能隔离已配置（资源配额、限流、熔断）\n""",
        "6": """\n\n### 6.0 SaaS 合规审查清单\n\n| 检查项 | 适用法规/标准 | 验证方式 | 责任人 |\n|:---|:---|:---|:---|\n| SOC2 Type II | AICPA信任服务标准 | 第三方审计报告 | 安全工程师 |\n| GDPR合规 | 欧盟通用数据保护条例 | 数据处理记录(DPR)+隐私政策 | 法务+DPO |\n| ISO27001 | 信息安全管理体系 | 认证机构审核 | 安全工程师 |\n| 数据驻留 | 客户合同/当地法规 | 数据中心部署文档 | 运维负责人 |\n| SLA保障 | 服务等级协议 | 可用性监控报告 | SRE负责人 |\n| 支付合规 | PCI DSS（如涉及支付） | QSA评估报告 | 安全工程师 |\n""",
        "8": """\n\n### 8.0 SaaS 里程碑评审闸门（Stage-Gate）\n\n| 阶段 | 评审标准 | 通过条件 | 参与角色 |\n|:---|:---|:---|:---|\n| **Gate 0: 立项** | PRD框架+商业模式验证 | 8章结构完整、TAM/SAM可估算 | PM、产品总监 |\n| **Gate 1: 设计** | 原型+技术架构方案 | UI评审通过、租户隔离方案确认 | PM、UI、架构师 |\n| **Gate 2: Alpha** | 核心功能+MVP租户 onboarding | 种子客户NPS≥30 | 开发、PM、客户成功 |\n| **Gate 3: Beta** | 付费转化+稳定性验证 | 付费转化率≥目标值、MTTR<4h | 开发、测试、销售 |\n| **Gate 4: GA发布** | 正式发布+规模化获客 | 系统可用性≥99.9%、支持流程就绪 | 运维、PM、市场 |\n| **Gate 5: 规模化** | 发布后90天评估 | NRR≥100%、CAC回收期<12月 | PM、增长、财务 |\n""",
    },
    "ecommerce": {
        "1": """\n\n### 1.0 电商产品 MVP 范围决策矩阵\n\n| 产品形态 | 适用场景 | 包含章节 | 裁剪建议 |\n|:---|:---|:---|:---|\n| **自营电商** | 品牌官方商城/D2C | 全部9章 | 第4章重点扩展供应链模块 |\n| **平台模式** | 多商家入驻（B2B2C） | 全部9章 | 第5章增加商家数据隔离、第6章增加平台责任 |\n| **社交电商** | 直播带货/社群团购 | 全部9章 | 第3章增加分享裂变流程、第7章增加社交指标 |\n| **O2O** | 同城配送/即时零售 | 全部9章 | 第3章增加LBS调度流程、第5章增加库存同步 |\n\n#### 电商关键指标数据来源\n\n| 占位符类型 | 建议数据来源 | 获取方式 |\n|:---|:---|:---|\n| GMV/订单量 | 交易系统统计 | 数据团队 |\n| 转化率 | 流量分析系统（GA/神策） | 增长团队 |\n| 客单价/复购率 | 用户交易记录聚合 | 数据团队 |\n| 退货率 | 售后系统统计 | 运营团队 |\n| 库存周转 | ERP/WMS系统 | 供应链团队 |\n\n#### RACI 矩阵模板\n\n| 事项 | 负责(R) | 批准(A) | 咨询(C) | 告知(I) |\n|:---|:---|:---|:---|:---|\n| PRD评审 | 产品经理 | 产品总监 | 技术负责人、运营负责人 | 客服、仓储 |\n| 范围变更 | 产品经理 | 产品总监 | 技术负责人 | 所有干系人 |\n| 支付合规审查 | 风控负责人 | 产品总监 | 法务、支付机构 | 财务 |\n| 大促上线决策 | 技术负责人 | 产品总监 | 运营、风控 | 全部门 |\n""",
        "2": """\n\n### 2.0 电商用户故事验收标准模板\n\n每个用户故事必须包含：\n- **触发条件**：什么场景下触发（如：用户浏览商品后、订单支付超时）\n- **前置条件**：系统/用户必须满足的状态（如：商品库存>0、用户已登录）\n- **基本流程**：标准操作步骤，含页面跳转和状态变更\n- **异常流程**：库存不足、支付失败、物流异常等处理方式\n- **验收标准（AC）**：可测试的通过/失败条件，需覆盖正向+异常流程\n- **优先级**：P0（必须）/ P1（应该）/ P2（可以）\n- **财务影响**：是否影响 GMV、是否涉及退款/佣金计算\n""",
        "5": """\n\n### 5.0 电商数据架构设计检查清单\n\n- [ ] 订单状态机已定义完整生命周期（待支付→已支付→待发货→已发货→已签收→已完成/售后）\n- [ ] 库存一致性策略已明确（预占/锁定/扣减时机，防止超卖）\n- [ ] 秒杀/大促峰值已规划限流和降级策略（库存预热、排队、熔断）\n- [ ] 商品SKU模型已支持多规格组合（颜色×尺码×版本）\n- [ ] 价格体系已支持多维度（原价/促销价/会员价/渠道价）\n- [ ] 对账数据已规划（支付流水、订单金额、佣金、退款）\n- [ ] 敏感字段已标注加密（收货地址、手机号、支付token）\n- [ ] 审计字段已覆盖全部资金操作（下单、支付、退款、结算）\n""",
        "6": """\n\n### 6.0 电商合规审查清单\n\n| 检查项 | 适用法规 | 验证方式 | 责任人 |\n|:---|:---|:---|:---|\n| 支付牌照/合规 | 非银行支付机构条例 | 支付服务商协议+PCI DSS | 风控负责人 |\n| 消费者权益保护 | 消费者权益保护法 | 退换货流程截图+客诉记录 | 运营负责人 |\n| 网络安全与数据 | 网络安全法、数据安全法 | 等保测评+数据分类分级 | 安全工程师 |\n| 税务合规 | 电子商务法、税收征管法 | 开票系统+税务申报记录 | 财务负责人 |\n| 商品信息合规 | 广告法、食品安全法（如适用） | 商品审核流程+抽检记录 | 运营负责人 |\n| 个人信息保护 | 个人信息保护法 | 隐私政策+用户授权记录 | 法务+DPO |\n""",
        "8": """\n\n### 8.0 电商里程碑评审闸门（Stage-Gate）\n\n| 阶段 | 评审标准 | 通过条件 | 参与角色 |\n|:---|:---|:---|:---|\n| **Gate 0: 立项** | PRD+商业模式+供应链方案 | 8章结构完整、毛利率模型可估算 | PM、产品总监 |\n| **Gate 1: 设计** | 原型+技术架构+支付集成方案 | UI评审通过、支付风控方案确认 | PM、UI、架构师 |\n| **Gate 2: 开发** | 核心交易链路完成 | 下单→支付→发货→售后全流程跑通 | 开发、测试 |\n| **Gate 3: 压测** | 全链路压测+资金安全验证 | 目标QPS达标、无资损漏洞 | 测试、风控、运维 |\n| **Gate 4: 上线** | 灰度发布+核心指标监控 | 支付成功率≥99.5%、无P0缺陷 | 运维、PM、运营 |\n| **Gate 5: 大促** | 大促全链路演练 | 容量≥峰值3倍、降级方案就绪 | 运维、风控、全部业务方 |\n""",
    },
}


@rate_limit(requests=30, window=60)
@router.post("", response_model=dict)
async def create_prd(
    data: PRDCreateRequest,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """Create a new PRD with template structure (AI generation is async)"""
    # Get project info for context
    result = await db.execute(
        select(Project).where(Project.id == data.project_id, Project.created_by == user_id)
    )
    project = result.scalar_one_or_none()

    template = data.template or "default"
    # Template takes priority for industry mapping; fallback to project.industry or general
    industry = TEMPLATE_TO_INDUSTRY.get(template)
    if not industry:
        industry = project.industry if project else "general"

    # Build empty chapters structure from template
    chapters = {}
    chapter_names = TEMPLATE_CHAPTERS.get(template, TEMPLATE_CHAPTERS["default"])
    framework = TEMPLATE_FRAMEWORK.get(template, TEMPLATE_FRAMEWORK["default"])
    for i, name in enumerate(chapter_names, 1):
        chapters[str(i)] = {
            "title": name,
            "content": "",
            "status": "draft"
        }

    # Build initial markdown with industry-specific framework infrastructure
    md_lines = [f"# {data.title}", ""]
    for i, name in enumerate(chapter_names, 1):
        md_lines.append(f"## {i}. {name}")
        md_lines.append("")
        framework_text = framework.get(str(i), "")
        if framework_text:
            md_lines.append(framework_text.strip())
        else:
            md_lines.append("（待填写）")
        md_lines.append("")

    content_struct = {
        "chapters": chapters,
        "template": template,
        "industry": industry,
    }

    new_prd = PRD(
        id=str(uuid.uuid4()),
        project_id=data.project_id,
        title=data.title,
        version="1.0",
        status=PRDStatus.DRAFT,
        content=content_struct,
        markdown="\n".join(md_lines),
        ai_generated={"template": template, "industry": industry, "ai_generated": False},
        created_by=user_id,
    )

    db.add(new_prd)
    await db.commit()
    await db.refresh(new_prd)

    return ResponseBuilder.success({
        "id": new_prd.id,
        "project_id": new_prd.project_id,
        "title": new_prd.title,
        "version": new_prd.version,
        "status": new_prd.status.value,
        "content": new_prd.content,
        "markdown": new_prd.markdown,
        "created_at": new_prd.created_at.isoformat() if new_prd.created_at else None,
        "updated_at": new_prd.updated_at.isoformat() if new_prd.updated_at else None,
    })


@rate_limit(requests=100, window=60)
@router.get("/{prd_id}", response_model=dict)
async def get_prd(
    prd_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """Get PRD by ID"""
    result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = result.scalar_one_or_none()

    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    return ResponseBuilder.success({
        "id": prd.id,
        "project_id": prd.project_id,
        "title": prd.title,
        "version": prd.version,
        "status": prd.status.value,
        "content": prd.content,
        "markdown": prd.markdown,
        "created_at": prd.created_at.isoformat() if prd.created_at else None,
        "updated_at": prd.updated_at.isoformat() if prd.updated_at else None,
    })


async def _run_compliance_check(prd_id: str, title: str, industry: str, markdown: str):
    """Background task: run compliance check after PRD update"""
    try:
        from app.agents.agents.compliance_checker import ComplianceChecker
        checker = ComplianceChecker()
        result = await checker.execute({
            "product_name": title,
            "industry": industry,
            "features": [markdown[:5000]],  # Truncate for analysis
        })
        if result.success:
            logging.info("Compliance check passed for PRD %s", prd_id)
        else:
            logging.warning("Compliance check found issues for PRD %s: %s", prd_id, result.error)
    except Exception:
        logging.exception("Compliance check failed for PRD %s", prd_id)


async def _index_prd_in_background(source_id: str, content: str, metadata: dict):
    """Background task: index PRD content with its own DB session"""
    from app.core.database import AsyncSessionLocal
    from app.services.memory_indexer import memory_indexer
    async with AsyncSessionLocal() as db:
        try:
            await memory_indexer.index_document(
                db=db,
                source_type="prd",
                source_id=source_id,
                content=content,
                metadata=metadata,
            )
        except Exception:
            logging.exception("Failed to index PRD %s into memory", source_id)


@rate_limit(requests=30, window=60)
@router.put("/{prd_id}", response_model=dict)
async def update_prd(
    prd_id: str,
    data: PRDUpdateRequest,
    background_tasks: BackgroundTasks,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """Update PRD"""
    result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = result.scalar_one_or_none()

    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    # Create version snapshot before update (if content changed)
    if data.markdown is not None or data.content is not None:
        # Debounce: only create version if last one is > 60s old
        last_ver_result = await db.execute(
            select(PRDVersion.created_at)
            .where(PRDVersion.prd_id == prd_id)
            .order_by(PRDVersion.version_number.desc())
            .limit(1)
        )
        last_ver_time = last_ver_result.scalar_one_or_none()
        now = datetime.now(timezone.utc)
        # SQLite returns naive datetimes; make aware for safe subtraction
        if last_ver_time is not None and last_ver_time.tzinfo is None:
            last_ver_time = last_ver_time.replace(tzinfo=timezone.utc)

        if not last_ver_time or (now - last_ver_time).total_seconds() > 60:
            version_count = await db.execute(
                select(func.count()).where(PRDVersion.prd_id == prd_id)
            )
            next_version = (version_count.scalar() or 0) + 1

            version = PRDVersion(
                prd_id=prd_id,
                version_number=next_version,
                title=prd.title,
                markdown=prd.markdown or "",
                content=json.dumps(prd.content) if prd.content else "",
                created_by=user_id,
            )
            db.add(version)

    if data.title is not None:
        prd.title = data.title
    if data.content is not None:
        prd.content = data.content
    if data.markdown is not None:
        prd.markdown = sanitize_html(data.markdown)
    if data.status is not None:
        try:
            prd.status = PRDStatus(data.status)
        except ValueError:
            raise AppException(f"Invalid status: {data.status}", code="INVALID_STATUS", status_code=400)

    await db.commit()
    await db.refresh(prd)

    # Trigger incremental compliance check when content changes
    if data.markdown is not None:
        industry = prd.content.get("industry", "general") if isinstance(prd.content, dict) else "general"
        background_tasks.add_task(
            _run_compliance_check,
            prd_id=prd.id,
            title=prd.title,
            industry=industry,
            markdown=prd.markdown or "",
        )

    # Index PRD content into semantic memory (background task with its own DB session)
    if data.markdown is not None:
        background_tasks.add_task(
            _index_prd_in_background,
            source_id=prd_id,
            content=data.markdown,
            metadata={"title": prd.title, "project_id": prd.project_id},
        )

    return ResponseBuilder.success({
        "id": prd.id,
        "project_id": prd.project_id,
        "title": prd.title,
        "version": prd.version,
        "status": prd.status.value,
        "content": prd.content,
        "markdown": prd.markdown,
        "created_at": prd.created_at.isoformat() if prd.created_at else None,
        "updated_at": prd.updated_at.isoformat() if prd.updated_at else None,
    })


@rate_limit(requests=20, window=60)
@router.delete("/{prd_id}", response_model=dict)
async def delete_prd(
    prd_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """Delete PRD"""
    result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = result.scalar_one_or_none()

    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    prd.soft_delete()
    await db.commit()

    return ResponseBuilder.success({
        "id": prd_id,
        "deleted": True
    })


@rate_limit(requests=10, window=60)
async def _summarize_chapter_in_background(
    prd_id: str, chapter_id: str, markdown: str, chapter_title: str
):
    """Background task: summarize a generated chapter for cross-chapter context."""
    from app.core.database import AsyncSessionLocal
    async with AsyncSessionLocal() as session:
        try:
            summary = await ai_service.summarize_chapter(markdown, chapter_title)
            if summary:
                result = await session.execute(
                    select(PRD).where(PRD.id == prd_id, PRD.deleted_at.is_(None))
                )
                prd_record = result.scalar_one_or_none()
                if prd_record:
                    content = prd_record.content or {}
                    content.setdefault("chapters", {}).setdefault(chapter_id, {})["summary"] = summary
                    prd_record.content = content
                    await session.commit()
        except Exception:
            logging.exception("Background summarization failed for PRD %s chapter %s", prd_id, chapter_id)


@router.post("/{prd_id}/generate-stream")
async def generate_prd_stream(
    prd_id: str,
    data: PRDGenerateRequest,
    background_tasks: BackgroundTasks,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """Stream PRD chapter generation using SSE"""
    result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = result.scalar_one_or_none()

    if not prd:
        async def error_stream():
            yield f"data: {json.dumps({'type': 'error', 'message': 'PRD not found'})}\n\n"
        return StreamingResponse(error_stream(), media_type="text/event-stream")

    industry = prd.content.get("industry", "general") if prd.content else "general"

    # Fetch project info, personas, and competitors concurrently
    project_context = {}
    if prd.project_id:
        import asyncio
        proj_task = db.execute(select(Project).where(Project.id == prd.project_id))
        personas_task = db.execute(select(Persona).where(Persona.project_id == prd.project_id, Persona.deleted_at.is_(None)))
        competitors_task = db.execute(select(Competitor).where(Competitor.project_id == prd.project_id, Competitor.deleted_at.is_(None)))
        proj_result, personas_result, competitors_result = await asyncio.gather(
            proj_task, personas_task, competitors_task
        )

        project = proj_result.scalar_one_or_none()
        if project:
            project_context = {
                "project_name": project.name,
                "project_description": project.description or "",
                "project_industry": project.industry or industry,
            }

        personas_list = personas_result.scalars().all()
        if personas_list:
            project_context["personas"] = [
                {
                    "name": p.name,
                    "role": p.role,
                    "description": p.description or "",
                }
                for p in personas_list
            ]

        competitors_list = competitors_result.scalars().all()
        if competitors_list:
            project_context["competitors"] = [
                {
                    "name": c.name,
                    "description": c.description or "",
                }
                for c in competitors_list
            ]

    # Resolve actual chapter title from the PRD's template structure
    chapter_title = None
    if prd.content and isinstance(prd.content, dict):
        chapters = prd.content.get("chapters", {})
        if data.chapter in chapters and isinstance(chapters[data.chapter], dict):
            chapter_title = chapters[data.chapter].get("title")

    # Merge PRD ai_generated metadata with project context
    ai_context = dict(prd.ai_generated) if prd.ai_generated else {}
    ai_context.update(project_context)

    # Build cross-chapter context from already-generated chapters for consistency.
    # For Chapter 1 always use raw content (it contains critical decision tables in markdown format).
    # For other chapters prefer AI-generated summary; fallback to raw content truncated to 3000 chars.
    existing_chapters: Dict[str, str] = {}
    if prd.content and isinstance(prd.content, dict):
        all_chapters = prd.content.get("chapters", {})
        for cid, cinfo in all_chapters.items():
            if cid != data.chapter and isinstance(cinfo, dict) and cinfo.get("status") == "generated":
                raw = cinfo.get("content", "")
                # Chapter 1 contains role tables, MVP boundary, metrics in markdown format.
                # Summary is plain text without table delimiters, so extraction methods fail.
                # Always use raw content for Chapter 1 to ensure structured data can be parsed.
                if cid == "1" and raw:
                    existing_chapters[cid] = raw[:4000] if len(raw) > 4000 else raw
                else:
                    summary = cinfo.get("summary", "")
                    if summary and len(summary.strip()) > 50:
                        existing_chapters[cid] = summary
                    else:
                        existing_chapters[cid] = raw[:3000] if len(raw) > 3000 else raw

    async def event_generator():
        import asyncio
        full_markdown = ""
        try:
            async for chunk in ai_service.generate_prd_chapter_stream(
                chapter=data.chapter,
                prompt=data.prompt,
                context=ai_context,
                industry=industry,
                chapter_title=chapter_title,
                existing_chapters=existing_chapters if existing_chapters else None,
                bypass_cache=data.bypass_cache,
            ):
                full_markdown += chunk
                yield f"data: {json.dumps({'type': 'chunk', 'text': chunk})}\n\n"
                await asyncio.sleep(0)  # Force event loop yield so data flushes to network

            # Save to DB after stream completes.
            # prd.content[chapters] = structured storage (summary, status, per-chapter content)
            # prd.markdown = canonical user-facing document, managed by the frontend via PUT /{id}
            # The frontend receives cleaned_markdown in the "done" SSE event and
            # persists it separately via prdApi.update(); this dual-write ensures
            # the structured backup survives even if the frontend save fails.
            cleaned_markdown = ai_service._clean_prd_output(full_markdown)
            content = prd.content or {"chapters": {}, "template": "standard", "industry": industry}
            content["chapters"] = content.get("chapters", {})
            content["chapters"][data.chapter] = {
                "title": chapter_title or f"Chapter {data.chapter}",
                "content": cleaned_markdown,
                "status": "generated"
            }

            prd.content = content
            await db.commit()

            yield f"data: {json.dumps({'type': 'done', 'markdown': cleaned_markdown})}\n\n"

            # Schedule summarization as a background task so it's not tied to
            # the SSE connection lifecycle (avoids GeneratorExit cancellation)
            background_tasks.add_task(
                _summarize_chapter_in_background,
                prd_id=prd_id,
                chapter_id=data.chapter,
                markdown=cleaned_markdown,
                chapter_title=chapter_title or f"Chapter {data.chapter}",
            )

        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ============== Document Export Helpers ==============

def _markdown_to_docx(markdown_text: str, title: str) -> bytes:
    """Convert markdown text to DOCX bytes."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    lines = markdown_text.splitlines()
    i = 0
    in_code_block = False
    code_lines: List[str] = []
    list_stack: List[int] = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.style = "Quote"
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove inline formatting markers for heading text
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_heading(text, level=min(level, 6))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.1)
            _add_formatted_runs(p, text)
            i += 1
            continue

        # Numbered list
        numbered_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered_match:
            indent = len(numbered_match.group(1))
            text = numbered_match.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.1)
            _add_formatted_runs(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_runs(p, line)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_formatted_runs(paragraph, text: str):
    """Add text runs to paragraph with bold/italic formatting."""
    # Split by bold markers first, preserving delimiters
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Check for italic within non-bold parts
            sub_parts = re.split(r"(\*.+?\*)", part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*") and not sub.startswith("**"):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub)


def _markdown_to_pdf(markdown_text: str, title: str) -> bytes:
    """Convert markdown text to PDF bytes."""
    from fpdf import FPDF

    class PRDPDF(FPDF):
        def header(self):
            self.set_font(self._font_name, "", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, title, align="C", ln=True)
            self.ln(2)

        def footer(self):
            self.set_y(-15)
            self.set_font(self._font_name, "", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"第 {self.page_no()} 页", align="C")

    pdf = PRDPDF()

    # Try to load a system CJK font for Chinese support, fallback to built-in
    _cjk_font_paths = [
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/System/Library/Fonts/PingFang.ttc",
    ]
    pdf._font_name = "Helvetica"
    for font_path in _cjk_font_paths:
        if os.path.exists(font_path):
            try:
                pdf.add_font("CJK", "", font_path, uni=True)
                pdf.add_font("CJK", "B", font_path, uni=True)
                pdf.add_font("CJK", "I", font_path, uni=True)
                pdf._font_name = "CJK"
                break
            except Exception:
                continue

    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font(pdf._font_name, "", 11)

    lines = markdown_text.splitlines()
    i = 0
    in_code_block = False
    code_buffer: List[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code_block:
                pdf.set_font("Courier", "", 9)
                pdf.set_fill_color(245, 245, 245)
                code_text = "\n".join(code_buffer)
                pdf.multi_cell(0, 5, code_text)
                pdf.ln(2)
                code_buffer = []
                pdf.set_font(pdf._font_name, "", 11)
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if not line.strip():
            pdf.ln(2)
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            sizes = {1: 18, 2: 14, 3: 12, 4: 11, 5: 10, 6: 10}
            size = sizes.get(level, 11)
            pdf.set_font(pdf._font_name, "B", size)
            pdf.set_text_color(33, 37, 41)
            pdf.multi_cell(0, 6, text)
            pdf.ln(1)
            pdf.set_font(pdf._font_name, "", 11)
            pdf.set_text_color(0, 0, 0)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(3)
            i += 1
            continue

        # Lists
        list_match = re.match(r"^(\s*)[-*+\d]+\.\s+(.+)$", line)
        is_list = list_match is not None

        # Determine indent and prefix
        prefix = ""
        indent = 0
        bullet_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        numbered_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            prefix = "• "
        elif numbered_match:
            indent = len(numbered_match.group(1))
            text = numbered_match.group(2)
            prefix = "◦ "
        else:
            text = line

        if is_list:
            pdf.set_x(15 + indent * 3)
            pdf.cell(5, 5, prefix)
            _write_pdf_formatted_text(pdf, text)
            pdf.ln(1)
        else:
            _write_pdf_formatted_text(pdf, text)
            pdf.ln(1)

        i += 1

    return pdf.output(dest="S")


def _write_pdf_formatted_text(pdf, text: str):
    """Write text to PDF with bold/italic support."""
    # Simple approach: render line by line switching fonts
    # Split by bold segments
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            pdf.set_font(pdf._font_name, "B", 11)
            pdf.write(5, inner)
            pdf.set_font(pdf._font_name, "", 11)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            inner = part[1:-1]
            pdf.set_font(pdf._font_name, "I", 11)
            pdf.write(5, inner)
            pdf.set_font(pdf._font_name, "", 11)
        else:
            pdf.write(5, part)


@rate_limit(requests=100, window=60)
@router.get("/{prd_id}/export", response_model=dict)
async def export_prd(
    prd_id: str,
    format: str = "markdown",
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """Export PRD in various formats: markdown, json, pdf, docx."""
    result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = result.scalar_one_or_none()

    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    if format == "json":
        content = json.dumps(prd.ai_generated, ensure_ascii=False, indent=2)
        return ResponseBuilder.success({
            "format": format,
            "content": content,
            "filename": f"prd_{prd.title or prd_id}.json"
        })

    markdown_content = prd.markdown or "# PRD Content"

    if format == "docx":
        try:
            binary = await asyncio.to_thread(_markdown_to_docx, markdown_content, prd.title or "PRD")
            b64 = base64.b64encode(binary).decode("utf-8")
            return ResponseBuilder.success({
                "format": format,
                "content": b64,
                "encoding": "base64",
                "filename": f"prd_{prd.title or prd_id}.docx"
            })
        except Exception as e:
            raise AppException(f"DOCX export failed: {str(e)}", code="EXPORT_ERROR", status_code=500)

    if format == "pdf":
        try:
            binary = await asyncio.to_thread(_markdown_to_pdf, markdown_content, prd.title or "PRD")
            b64 = base64.b64encode(binary).decode("utf-8")
            return ResponseBuilder.success({
                "format": format,
                "content": b64,
                "encoding": "base64",
                "filename": f"prd_{prd.title or prd_id}.pdf"
            })
        except Exception as e:
            raise AppException(f"PDF export failed: {str(e)}", code="EXPORT_ERROR", status_code=500)

    # Default: markdown
    return ResponseBuilder.success({
        "format": format,
        "content": markdown_content,
        "filename": f"prd_{prd.title or prd_id}.md"
    })


# ============== Version History ==============

@rate_limit(requests=100, window=60)
@router.get("/{prd_id}/versions", response_model=dict)
async def list_prd_versions(
    prd_id: str,
    limit: int = 20,
    offset: int = 0,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """List PRD version history"""
    # Verify PRD ownership
    prd_result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = prd_result.scalar_one_or_none()
    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    query = select(PRDVersion).where(PRDVersion.prd_id == prd_id).order_by(desc(PRDVersion.created_at))
    total_result = await db.execute(select(func.count()).select_from(query.subquery()))
    total = total_result.scalar() or 0

    query = query.offset(offset).limit(limit)
    result = await db.execute(query)
    versions = result.scalars().all()

    return ResponseBuilder.paginated(
        data=[{
            "id": v.id,
            "version_number": v.version_number,
            "title": v.title,
            "change_summary": v.change_summary,
            "created_at": v.created_at.isoformat() if v.created_at else None,
        } for v in versions],
        page=offset // limit + 1 if limit > 0 else 1,
        limit=limit,
        total=total,
    )


@rate_limit(requests=100, window=60)
@router.get("/{prd_id}/versions/{version_id}", response_model=dict)
async def get_prd_version(
    prd_id: str,
    version_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """Get a specific PRD version"""
    prd_result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = prd_result.scalar_one_or_none()
    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    result = await db.execute(
        select(PRDVersion).where(PRDVersion.id == version_id, PRDVersion.prd_id == prd_id)
    )
    version = result.scalar_one_or_none()
    if not version:
        raise AppException("Version not found", code="NOT_FOUND", status_code=404)

    return ResponseBuilder.success({
        "id": version.id,
        "version_number": version.version_number,
        "title": version.title,
        "markdown": version.markdown,
        "content": version.content,
        "created_at": version.created_at.isoformat() if version.created_at else None,
    })


@rate_limit(requests=30, window=60)
@router.post("/{prd_id}/versions/{version_id}/restore", response_model=dict)
async def restore_prd_version(
    prd_id: str,
    version_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    """Restore PRD to a specific version"""
    prd_result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = prd_result.scalar_one_or_none()
    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    result = await db.execute(
        select(PRDVersion).where(PRDVersion.id == version_id, PRDVersion.prd_id == prd_id)
    )
    version = result.scalar_one_or_none()
    if not version:
        raise AppException("Version not found", code="NOT_FOUND", status_code=404)

    # Create a new version snapshot of current state before restoring
    version_count = await db.execute(
        select(func.count()).where(PRDVersion.prd_id == prd_id)
    )
    next_version = (version_count.scalar() or 0) + 1

    backup = PRDVersion(
        prd_id=prd_id,
        version_number=next_version,
        title=prd.title,
        markdown=prd.markdown or "",
        content=json.dumps(prd.content) if prd.content else "",
        change_summary="Auto-backup before restore",
        created_by=user_id,
    )
    db.add(backup)

    # Restore
    prd.title = version.title
    prd.markdown = version.markdown
    try:
        prd.content = json.loads(version.content) if version.content else {}
    except Exception:
        prd.content = {}

    await db.commit()
    await db.refresh(prd)

    return ResponseBuilder.success({
        "message": f"Restored to version {version.version_number}",
        "prd_id": prd.id,
        "restored_version": version.version_number,
    })


@rate_limit(requests=100, window=60)
@router.get("/{prd_id}/generation-progress", response_model=dict)
async def get_generation_progress(
    prd_id: str,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """Get PRD chapter generation progress"""
    result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = result.scalar_one_or_none()

    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    chapters = {}
    if prd.content and isinstance(prd.content, dict):
        chapters = prd.content.get("chapters", {})

    total = len(chapters)
    generated = sum(1 for c in chapters.values() if isinstance(c, dict) and c.get("status") == "generated")
    pending = total - generated

    progress = {
        "total_chapters": total,
        "generated": generated,
        "pending": pending,
        "percentage": round(generated / total * 100, 1) if total else 0,
        "chapters": [
            {
                "id": cid,
                "title": c.get("title", f"Chapter {cid}") if isinstance(c, dict) else f"Chapter {cid}",
                "status": c.get("status", "draft") if isinstance(c, dict) else "draft",
            }
            for cid, c in chapters.items()
        ],
    }

    return ResponseBuilder.success(progress)


@rate_limit(requests=100, window=60)
@router.get("/{prd_id}/diff", response_model=dict)
async def diff_prd_versions(
    prd_id: str,
    from_version: int = Query(..., description="Source version number"),
    to_version: int = Query(..., description="Target version number"),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """Compare two PRD versions and return unified diff"""
    # Verify PRD ownership
    prd_result = await db.execute(
        select(PRD).where(PRD.id == prd_id, PRD.created_by == user_id, PRD.deleted_at.is_(None))
    )
    prd = prd_result.scalar_one_or_none()
    if not prd:
        raise AppException("PRD not found", code="NOT_FOUND", status_code=404)

    # Fetch both versions
    from_result = await db.execute(
        select(PRDVersion).where(
            PRDVersion.prd_id == prd_id,
            PRDVersion.version_number == from_version
        )
    )
    to_result = await db.execute(
        select(PRDVersion).where(
            PRDVersion.prd_id == prd_id,
            PRDVersion.version_number == to_version
        )
    )
    from_v = from_result.scalar_one_or_none()
    to_v = to_result.scalar_one_or_none()

    if not from_v or not to_v:
        missing = []
        if not from_v:
            missing.append(str(from_version))
        if not to_v:
            missing.append(str(to_version))
        raise AppException(f"Version(s) not found: {', '.join(missing)}", code="NOT_FOUND", status_code=400)

    from_lines = (from_v.markdown or "").splitlines(keepends=True)
    to_lines = (to_v.markdown or "").splitlines(keepends=True)

    diff = list(difflib.unified_diff(
        from_lines,
        to_lines,
        fromfile=f"v{from_version}",
        tofile=f"v{to_version}",
        lineterm="",
    ))

    # Also compute a high-level summary
    added = sum(1 for line in diff if line.startswith("+"))
    removed = sum(1 for line in diff if line.startswith("-"))
    unchanged = sum(1 for line in diff if line.startswith(" "))

    return ResponseBuilder.success({
        "from_version": from_version,
        "to_version": to_version,
        "from_title": from_v.title,
        "to_title": to_v.title,
        "from_created_at": from_v.created_at.isoformat() if from_v.created_at else None,
        "to_created_at": to_v.created_at.isoformat() if to_v.created_at else None,
        "diff": "".join(diff),
        "summary": {
            "lines_added": added,
            "lines_removed": removed,
            "lines_unchanged": unchanged,
            "total_changes": added + removed,
        },
    })